"""
Stage 7 – Report Generation
Produces JSON, CSV, Excel workbook, and Word (.docx) document.
Domain name and labels are driven by PipelineConfig — fully generic.
"""

import json
import csv
import os
from datetime import date

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, Reference

try:
    from docx import Document
    from docx.shared import Pt, RGBColor as DocxRGB, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from .config import PipelineConfig, DEFAULT_FMCG_CONFIG

# Colour palette
NAVY      = "1B3A5C"
TEAL      = "0D7377"
GOLD      = "C9962B"
LIGHT_BG  = "EFF4FA"
MID_BG    = "D6E4F0"
WHITE     = "FFFFFF"
DARK_TEXT = "1A1A2E"
GREY_TEXT = "5A6375"
GREEN_OK  = "1E8449"

DEAL_TYPE_COLORS = {
    "Acquisition":       "1B3A5C",
    "M&A":               "1B3A5C",
    "Merger":            "0D7377",
    "Investment":        "1E8449",
    "Divestiture":       "7D3C98",
    "Stake Acquisition": "1A5276",
    "Joint Venture":     "935116",
    "Other":             "5A6375",
}


def _hex_fill(hex_color: str) -> PatternFill:
    return PatternFill("solid", fgColor=hex_color)


def _font(bold=False, color=DARK_TEXT, size=10, italic=False) -> Font:
    return Font(bold=bold, color=color, size=size, italic=italic, name="Calibri")


def _border() -> Border:
    s = Side(style="thin")
    return Border(left=s, right=s, top=s, bottom=s)


def _align(h="left", v="center", wrap=False) -> Alignment:
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)


def _w(ws, row, col, value, bold=False, color=DARK_TEXT, size=10,
       bg=None, align_h="left", wrap=False, italic=False, border=False):
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = _font(bold=bold, color=color, size=size, italic=italic)
    cell.alignment = _align(h=align_h, wrap=wrap)
    if bg:
        cell.fill = _hex_fill(bg)
    if border:
        cell.border = _border()
    return cell


def _build_newsletter_sheet(ws, articles: list[dict], pipeline_log: list[dict], config: PipelineConfig):
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 3
    for col_letter in ["B", "C", "D", "E", "F", "G", "H"]:
        ws.column_dimensions[col_letter].width = 22
    ws.column_dimensions["I"].width = 3

    issue_date = date.today().strftime("%B %d, %Y")
    top_articles = sorted(articles, key=lambda a: a.get("relevance_score", 0), reverse=True)[:20]

    row = 1

    # Banner
    ws.merge_cells(f"B{row}:H{row}")
    _w(ws, row, 2, config.domain_name.upper(), bold=True, color=WHITE,
       size=20, bg=NAVY, align_h="center")
    ws.row_dimensions[row].height = 36
    row += 1

    ws.merge_cells(f"B{row}:H{row}")
    _w(ws, row, 2, f"Intelligence Report  |  {issue_date}",
       color=WHITE, size=11, bg=TEAL, align_h="center", italic=True)
    ws.row_dimensions[row].height = 22
    row += 1

    ws.merge_cells(f"B{row}:H{row}")
    _w(ws, row, 2, "", bg=MID_BG)
    ws.row_dimensions[row].height = 8
    row += 1

    # KPI strip
    deal_types: dict[str, int] = {}
    for a in top_articles:
        dt = a.get("deal_type_detected", "Other")
        deal_types[dt] = deal_types.get(dt, 0) + 1

    acq  = deal_types.get("Acquisition", 0) + deal_types.get("M&A", 0) + deal_types.get("Merger", 0)
    inv  = deal_types.get("Investment", 0)
    div  = deal_types.get("Divestiture", 0)

    kpis = [
        ("TOTAL RECORDS", str(len(top_articles)), NAVY),
        ("ACQUISITIONS",  str(acq),  TEAL),
        ("INVESTMENTS",   str(inv),  GREEN_OK),
        ("DIVESTITURES",  str(div),  "7D3C98"),
    ]
    kpi_cols = [2, 4, 6, 8]
    for (label, val, color), col in zip(kpis, kpi_cols):
        ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col + 1)
        _w(ws, row, col, label, bold=True, color=WHITE, size=9, bg=color, align_h="center")
    ws.row_dimensions[row].height = 18
    row += 1

    for (label, val, color), col in zip(kpis, kpi_cols):
        ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col + 1)
        _w(ws, row, col, val, bold=True, color=color, size=22, align_h="center", bg=LIGHT_BG)
    ws.row_dimensions[row].height = 34
    row += 1

    ws.merge_cells(f"B{row}:H{row}")
    _w(ws, row, 2, "", bg=MID_BG)
    ws.row_dimensions[row].height = 10
    row += 1

    # Articles table
    ws.merge_cells(f"B{row}:H{row}")
    _w(ws, row, 2, "  TOP RECORDS BY RELEVANCE SCORE",
       bold=True, color=WHITE, size=10, bg=NAVY)
    ws.row_dimensions[row].height = 20
    row += 1

    headers = ["Date", "Title", "Summary", "Source", "Deal Type", "Score"]
    col_defs = [(2, 1), (3, 2), (5, 2), (7, 1), (8, 1), (9, 1)]
    for header, (col, span) in zip(headers, col_defs):
        end_col = col + span - 1
        if span > 1:
            ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=end_col)
        _w(ws, row, col, header, bold=True, color=WHITE, size=9, bg=TEAL, align_h="center", border=True)
    ws.row_dimensions[row].height = 18
    row += 1

    for i, article in enumerate(top_articles):
        bg = WHITE if i % 2 == 0 else LIGHT_BG
        pub_date = article.get("published_date", "")
        title    = article.get("title", "")
        url      = article.get("url", "")
        summary  = (article.get("summary", "") or "")[:200]
        if len(article.get("summary", "") or "") > 200:
            summary += "…"
        source   = article.get("source", "")
        deal_t   = article.get("deal_type_detected", "Other")
        score    = f"{article.get('relevance_score', 0):.1f}"
        deal_bg  = DEAL_TYPE_COLORS.get(deal_t, GREY_TEXT)

        _w(ws, row, 2, pub_date, size=9, bg=bg, align_h="center", border=True)
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
        # Title as clickable hyperlink
        title_cell = ws.cell(row=row, column=3, value=title)
        if url and url.startswith("http"):
            title_cell.hyperlink = url
            title_cell.font = Font(bold=True, color="0563C1", underline="single", size=9, name="Calibri")
        else:
            title_cell.font = _font(bold=True, size=9)
        title_cell.fill = _hex_fill(bg)
        title_cell.alignment = _align(wrap=True)
        title_cell.border = _border()
        ws.merge_cells(start_row=row, start_column=5, end_row=row, end_column=6)
        _w(ws, row, 5, summary, size=8, bg=bg, wrap=True, color=GREY_TEXT, border=True)
        _w(ws, row, 7, source, size=9, bg=bg, align_h="center", border=True)
        _w(ws, row, 8, deal_t, size=9, bg=deal_bg, color=WHITE, align_h="center", border=True)
        _w(ws, row, 9, score, size=9, bg=bg, align_h="center", border=True)
        ws.row_dimensions[row].height = 52
        row += 1

    ws.merge_cells(f"B{row}:H{row}")
    _w(ws, row, 2,
       f"Generated by {config.domain_name} Pipeline  |  {issue_date}  |  For internal use only",
       size=8, color=WHITE, bg=NAVY, align_h="center", italic=True, wrap=True)
    ws.row_dimensions[row].height = 28


def _build_summary_sheet(ws, articles: list[dict]):
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 3
    for col, w in {"B": 25, "C": 15, "D": 18, "E": 20, "F": 18, "G": 3}.items():
        ws.column_dimensions[col].width = w

    row = 1
    ws.merge_cells(f"B{row}:F{row}")
    _w(ws, row, 2, "SUMMARY — BY CATEGORY", bold=True, color=WHITE, size=13, bg=NAVY, align_h="center")
    ws.row_dimensions[row].height = 28
    row += 2

    groups: dict[str, list[dict]] = {}
    for a in articles:
        dt = a.get("deal_type_detected", "Other")
        groups.setdefault(dt, []).append(a)

    for i, h in enumerate(["Category", "# Records", "Avg Relevance", "Avg Credibility", "Sources (sample)"], start=2):
        _w(ws, row, i, h, bold=True, color=WHITE, size=10, bg=TEAL, align_h="center", border=True)
    ws.row_dimensions[row].height = 20
    row += 1

    chart_categories, chart_counts = [], []

    for dt, items in sorted(groups.items(), key=lambda x: -len(x[1])):
        avg_rel  = sum(a.get("relevance_score", 0) for a in items) / len(items)
        avg_cred = sum(a.get("credibility_score", 0) for a in items) / len(items)
        sources  = ", ".join(sorted({a.get("source", "") for a in items})[:3])
        bg_hex   = DEAL_TYPE_COLORS.get(dt, GREY_TEXT)

        _w(ws, row, 2, dt,           bold=True, color=WHITE, size=10, bg=bg_hex, border=True)
        _w(ws, row, 3, len(items),   size=10, align_h="center", border=True, bg=LIGHT_BG)
        _w(ws, row, 4, f"{avg_rel:.1f}",  size=10, align_h="center", border=True, bg=LIGHT_BG)
        _w(ws, row, 5, f"{avg_cred:.1f}", size=10, align_h="center", border=True, bg=LIGHT_BG)
        _w(ws, row, 6, sources,      size=9, color=GREY_TEXT, border=True, bg=LIGHT_BG, wrap=True)
        chart_categories.append(dt)
        chart_counts.append(len(items))
        ws.row_dimensions[row].height = 22
        row += 1

    row += 2
    chart = BarChart()
    chart.type = "col"
    chart.title = "Record Count by Category"
    chart.style = 10
    chart.y_axis.title = "Count"
    chart.width = 18
    chart.height = 12

    ws.cell(row=row, column=2, value="Category")
    ws.cell(row=row, column=3, value="Count")
    for i, (cat, cnt) in enumerate(zip(chart_categories, chart_counts), start=1):
        ws.cell(row=row + i, column=2, value=cat)
        ws.cell(row=row + i, column=3, value=cnt)

    data = Reference(ws, min_col=3, min_row=row, max_row=row + len(chart_categories))
    cats = Reference(ws, min_col=2, min_row=row + 1, max_row=row + len(chart_categories))
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    ws.add_chart(chart, f"B{row + len(chart_categories) + 2}")


def _hyperlink_cell(ws, row: int, col: int, url: str, display: str, bg: str):
    """Write a clickable hyperlink into an Excel cell."""
    cell = ws.cell(row=row, column=col, value=display)
    if url and url.startswith("http"):
        cell.hyperlink = url
        cell.font = Font(color="0563C1", underline="single", size=9, name="Calibri")
    else:
        cell.font = _font(size=9, color=GREY_TEXT)
    cell.fill = _hex_fill(bg)
    cell.alignment = _align(wrap=True)
    cell.border = _border()
    return cell


def _build_articles_sheet(ws, articles: list[dict]):
    ws.sheet_view.showGridLines = True
    headers = [
        "ID", "Title", "Source", "Published Date", "Category",
        "Domain Score", "Deal Score", "Recency Score", "Credibility Score",
        "Relevance Score", "Source Tier", "Credibility Flag", "Link Valid", "URL",
    ]
    field_map = [
        "id", "title", "source", "published_date", "deal_type_detected",
        "score_domain", "score_deal", "score_recency", "credibility_score",
        "relevance_score", "source_tier", "credibility_flag", "link_valid", "url",
    ]
    col_widths = [6, 50, 18, 14, 18, 13, 12, 13, 17, 15, 12, 30, 10, 40]

    for i, (h, w) in enumerate(zip(headers, col_widths), start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
        cell = ws.cell(row=1, column=i, value=h)
        cell.font = _font(bold=True, color=WHITE)
        cell.fill = _hex_fill(NAVY)
        cell.alignment = _align(h="center")
        cell.border = _border()
    ws.row_dimensions[1].height = 20

    # URL column index (1-based)
    url_col = len(field_map)

    for r, article in enumerate(articles, start=2):
        bg = LIGHT_BG if r % 2 == 0 else WHITE
        for c, field in enumerate(field_map, start=1):
            if field == "url":
                url_val = article.get("url", "")
                _hyperlink_cell(ws, r, c, url_val, url_val[:60] if url_val else "", bg)
            elif field == "link_valid":
                lv = article.get("link_valid")
                val = "✓" if lv else ("✗" if lv is False else "?")
                color = GREEN_OK if lv else ("C0392B" if lv is False else GREY_TEXT)
                cell = ws.cell(row=r, column=c, value=val)
                cell.font = _font(bold=True, color=color, size=9)
                cell.fill = _hex_fill(bg)
                cell.alignment = _align(h="center")
                cell.border = _border()
            else:
                val = article.get(field, "")
                cell = ws.cell(row=r, column=c, value=val)
                cell.font = _font(size=9)
                cell.fill = _hex_fill(bg)
                cell.alignment = _align(wrap=True)
                cell.border = _border()
        ws.row_dimensions[r].height = 30


def _build_pipeline_log_sheet(ws, pipeline_log: list[dict]):
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 3
    for col, w in {"B": 30, "C": 20, "D": 20, "E": 40, "F": 3}.items():
        ws.column_dimensions[col].width = w

    row = 1
    ws.merge_cells(f"B{row}:E{row}")
    _w(ws, row, 2, "PIPELINE EXECUTION LOG", bold=True, color=WHITE, size=13, bg=NAVY, align_h="center")
    ws.row_dimensions[row].height = 26
    row += 2

    for i, h in enumerate(["Stage", "Input Records", "Output Records", "Notes"], start=2):
        _w(ws, row, i, h, bold=True, color=WHITE, size=10, bg=TEAL, align_h="center", border=True)
    ws.row_dimensions[row].height = 20
    row += 1

    stage_colors = [NAVY, TEAL, GREEN_OK, "7D3C98", GOLD]
    for i, entry in enumerate(pipeline_log):
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        color = stage_colors[i % len(stage_colors)]
        _w(ws, row, 2, entry.get("stage", ""),  bold=True, color=WHITE, bg=color, size=10, border=True)
        _w(ws, row, 3, entry.get("input", ""),  size=10, align_h="center", bg=bg, border=True)
        _w(ws, row, 4, entry.get("output", ""), size=10, align_h="center", bg=bg, border=True)
        _w(ws, row, 5, entry.get("notes", ""),  size=9, color=GREY_TEXT, bg=bg, border=True, wrap=True)
        ws.row_dimensions[row].height = 28
        row += 1


def _build_docx(
    articles: list[dict],
    pipeline_log: list[dict],
    config: PipelineConfig,
) -> "Document":
    """Build a professional Word newsletter document."""
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    issue_date = date.today().strftime("%B %d, %Y")
    top_articles = sorted(articles, key=lambda a: a.get("relevance_score", 0), reverse=True)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _set_run(run, bold=False, size=11, rgb=None, italic=False):
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if rgb:
            run.font.color.rgb = DocxRGB(*rgb)

    def _heading(text, level=1, bold=True, size=14, rgb=(27, 58, 92), align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        _set_run(run, bold=bold, size=size, rgb=rgb)
        return p

    def _divider():
        p = doc.add_paragraph("─" * 80)
        p.runs[0].font.color.rgb = DocxRGB(200, 210, 225)
        p.runs[0].font.size = Pt(7)
        return p

    # ── Title banner ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title_p.add_run(config.domain_name.upper())
    _set_run(r1, bold=True, size=22, rgb=(27, 58, 92))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run(f"Intelligence Report  |  {issue_date}")
    _set_run(r2, italic=True, size=11, rgb=(13, 115, 119))

    _divider()

    # ── KPI summary ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading("  KEY METRICS", level=2, size=12, rgb=(13, 115, 119))

    deal_types: dict[str, int] = {}
    for a in top_articles:
        dt = a.get("deal_type_detected", "Other")
        deal_types[dt] = deal_types.get(dt, 0) + 1

    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.style = "Table Grid"
    kpi_labels = ["TOTAL DEALS", "ACQUISITIONS", "INVESTMENTS", "DIVESTITURES"]
    kpi_values = [
        str(len(top_articles)),
        str(deal_types.get("Acquisition", 0) + deal_types.get("M&A", 0) + deal_types.get("Merger", 0)),
        str(deal_types.get("Investment", 0)),
        str(deal_types.get("Divestiture", 0)),
    ]
    for col, (lbl, val) in enumerate(zip(kpi_labels, kpi_values)):
        lbl_cell = kpi_table.cell(0, col)
        lbl_cell.text = ""
        r = lbl_cell.paragraphs[0].add_run(lbl)
        _set_run(r, bold=True, size=8, rgb=(27, 58, 92))
        lbl_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        val_cell = kpi_table.cell(1, col)
        val_cell.text = ""
        r2 = val_cell.paragraphs[0].add_run(val)
        _set_run(r2, bold=True, size=20, rgb=(13, 115, 119))
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _divider()

    # ── Key highlights ────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading("  TOP 3 KEY HIGHLIGHTS", level=2, size=13, rgb=(27, 58, 92))
    doc.add_paragraph()

    rank_labels = ["🥇 #1", "🥈 #2", "🥉 #3"]
    for i, article in enumerate(top_articles[:3]):
        rank_p = doc.add_paragraph()
        r_rank = rank_p.add_run(f"{rank_labels[i]}  ")
        _set_run(r_rank, bold=True, size=11, rgb=(201, 150, 43))

        r_title = rank_p.add_run(article.get("title", ""))
        _set_run(r_title, bold=True, size=11, rgb=(27, 58, 92))

        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.left_indent = Cm(0.5)
        r_meta = meta_p.add_run(
            f"{article.get('source', '')}  ·  {article.get('published_date', '')}  "
            f"·  {article.get('deal_type_detected', '')}  "
            f"·  Score: {article.get('relevance_score', 0):.1f}"
        )
        _set_run(r_meta, size=9, rgb=(90, 99, 117), italic=True)

        summary_p = doc.add_paragraph()
        summary_p.paragraph_format.left_indent = Cm(0.5)
        r_sum = summary_p.add_run(article.get("summary", "")[:300])
        _set_run(r_sum, size=10, rgb=(50, 60, 80))

        # Clickable source URL
        url_val = article.get("url", "")
        if url_val and url_val.startswith("http"):
            url_p = doc.add_paragraph()
            url_p.paragraph_format.left_indent = Cm(0.5)
            r_url = url_p.add_run(f"Source: {url_val[:100]}")
            _set_run(r_url, size=8, rgb=(5, 99, 193), italic=True)
            # Word hyperlink via relationship
            try:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                from lxml import etree
                r_url.clear()
                hyperlink = OxmlElement("w:hyperlink")
                r_id = url_p.part.relate_to(
                    url_val,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                hyperlink.set(qn("r:id"), r_id)
                new_run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                color_el = OxmlElement("w:color")
                color_el.set(qn("w:val"), "0563C1")
                u_el = OxmlElement("w:u")
                u_el.set(qn("w:val"), "single")
                sz_el = OxmlElement("w:sz")
                sz_el.set(qn("w:val"), "16")
                rPr.append(color_el)
                rPr.append(u_el)
                rPr.append(sz_el)
                new_run.append(rPr)
                t_el = OxmlElement("w:t")
                t_el.text = f"Source: {url_val[:100]}"
                new_run.append(t_el)
                hyperlink.append(new_run)
                url_p._p.clear()
                url_p._p.append(hyperlink)
            except Exception:
                pass  # Fallback: plain text URL already set above

        if i < 2:
            doc.add_paragraph()

    _divider()

    # ── All deals table ───────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading(f"  ALL DEALS  ({len(top_articles)} records)", level=2, size=13, rgb=(27, 58, 92))
    doc.add_paragraph()

    deals_table = doc.add_table(rows=1, cols=6)
    deals_table.style = "Table Grid"
    deals_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_texts = ["Date", "Title", "Source", "Deal Type", "Score", "URL"]
    header_widths = [Cm(2.2), Cm(5.5), Cm(2.8), Cm(2.8), Cm(1.5), Cm(4.5)]
    for col_idx, (hdr, width) in enumerate(zip(header_texts, header_widths)):
        cell = deals_table.cell(0, col_idx)
        cell.width = width
        cell.text = ""
        r = cell.paragraphs[0].add_run(hdr)
        _set_run(r, bold=True, size=9, rgb=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for article in top_articles:
        row_cells = deals_table.add_row().cells
        values = [
            article.get("published_date", ""),
            article.get("title", ""),
            article.get("source", ""),
            article.get("deal_type_detected", "Other"),
            f"{article.get('relevance_score', 0):.1f}",
            article.get("url", ""),
        ]
        for col_idx, val in enumerate(values):
            row_cells[col_idx].text = ""
            if col_idx == 5 and val and val.startswith("http"):
                # Clickable hyperlink in URL column
                r = row_cells[col_idx].paragraphs[0].add_run(val[:60])
                _set_run(r, size=7, rgb=(5, 99, 193))
                try:
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    r.clear()
                    hyperlink = OxmlElement("w:hyperlink")
                    r_id = row_cells[col_idx].paragraphs[0].part.relate_to(
                        val,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                        is_external=True,
                    )
                    hyperlink.set(qn("r:id"), r_id)
                    new_run = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    color_el = OxmlElement("w:color")
                    color_el.set(qn("w:val"), "0563C1")
                    u_el = OxmlElement("w:u")
                    u_el.set(qn("w:val"), "single")
                    sz_el = OxmlElement("w:sz")
                    sz_el.set(qn("w:val"), "14")
                    rPr.extend([color_el, u_el, sz_el])
                    new_run.append(rPr)
                    t_el = OxmlElement("w:t")
                    t_el.text = val[:60]
                    new_run.append(t_el)
                    hyperlink.append(new_run)
                    row_cells[col_idx].paragraphs[0]._p.clear()
                    row_cells[col_idx].paragraphs[0]._p.append(hyperlink)
                except Exception:
                    pass
            else:
                r = row_cells[col_idx].paragraphs[0].add_run(str(val))
                _set_run(r, size=8, rgb=(26, 26, 46))
            row_cells[col_idx].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            )

    _divider()

    # ── Pipeline log ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading("  PIPELINE EXECUTION LOG", level=2, size=13, rgb=(27, 58, 92))
    doc.add_paragraph()

    for entry in pipeline_log:
        log_p = doc.add_paragraph()
        r_stage = log_p.add_run(f"[{entry.get('stage', '')}]  ")
        _set_run(r_stage, bold=True, size=10, rgb=(13, 115, 119))
        r_io = log_p.add_run(f"{entry.get('input', '')} → {entry.get('output', '')}  ")
        _set_run(r_io, bold=True, size=10, rgb=(27, 58, 92))
        r_notes = log_p.add_run(entry.get("notes", ""))
        _set_run(r_notes, size=9, rgb=(90, 99, 117), italic=True)

    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Generated by {config.domain_name} Pipeline  ·  {issue_date}  ·  For internal use only"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(footer_p.runs[0], size=8, rgb=(90, 99, 117), italic=True)

    return doc


def generate(
    articles: list[dict],
    pipeline_log: list[dict],
    output_dir: str = "output",
    config: PipelineConfig | None = None,
    progress_cb=None,
) -> dict:
    if config is None:
        config = DEFAULT_FMCG_CONFIG

    os.makedirs(output_dir, exist_ok=True)
    safe_name = config.domain_name.replace(" ", "_").replace("/", "-")

    # JSON
    json_path = os.path.join(output_dir, "processed_articles.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(articles, f, indent=2, default=str)

    # CSV
    csv_path = os.path.join(output_dir, "processed_articles.csv")
    if articles:
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=list(articles[0].keys()))
            writer.writeheader()
            writer.writerows(articles)

    # Excel
    xlsx_path = os.path.join(output_dir, f"{safe_name}_Report.xlsx")
    wb = openpyxl.Workbook()
    ws_news     = wb.active
    ws_news.title = "Report"
    ws_summary  = wb.create_sheet("Summary")
    ws_articles = wb.create_sheet("All Records")
    ws_log      = wb.create_sheet("Pipeline Log")

    _build_newsletter_sheet(ws_news, articles, pipeline_log, config)
    _build_summary_sheet(ws_summary, articles)
    _build_articles_sheet(ws_articles, articles)
    _build_pipeline_log_sheet(ws_log, pipeline_log)

    ws_news.sheet_properties.tabColor     = NAVY
    ws_summary.sheet_properties.tabColor  = TEAL
    ws_articles.sheet_properties.tabColor = GREEN_OK
    ws_log.sheet_properties.tabColor      = GOLD

    wb.save(xlsx_path)

    # Word (.docx)
    docx_path = os.path.join(output_dir, f"{safe_name}_Newsletter.docx")
    if _DOCX_AVAILABLE:
        try:
            doc = _build_docx(articles, pipeline_log, config)
            doc.save(docx_path)
        except Exception as exc:
            print(f"[Newsletter] Warning: docx generation failed: {exc}")
            docx_path = None
    else:
        print("[Newsletter] python-docx not installed — skipping .docx output")
        docx_path = None

    output_files = {"json": json_path, "csv": csv_path, "xlsx": xlsx_path}
    if docx_path:
        output_files["docx"] = docx_path

    fmt_list = ", ".join(k.upper() for k in output_files)
    msg = f"Generated {len(articles)} records → {fmt_list}"
    print(f"[Newsletter] {msg}")
    if progress_cb:
        progress_cb("newsletter", len(articles), len(articles), msg)

    return output_files
